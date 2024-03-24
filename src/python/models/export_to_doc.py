from docx import Document


def export_to_docx(data, file_name):
    doc = Document()
    for feature_name, feature_data in data.items():
        doc.add_heading(feature_name, level=1)
        doc.add_paragraph(str(feature_data))
    doc.save(file_name)
